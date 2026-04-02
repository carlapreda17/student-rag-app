import os
import uuid
import io
import pdfplumber
from docx import Document as DocxDocument
from fastapi import UploadFile, File, Form, HTTPException
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from qdrant_client import QdrantClient
from qdrant_client.http import models


load_dotenv()

# ── Qdrant ──────────────────────────────────────────────
qdrant = QdrantClient(url=os.getenv("QDRANT_URL", "http://localhost:6333"))
COLLECTION_NAME = "documente_studenti"

try:
    qdrant.get_collection(collection_name=COLLECTION_NAME)
except Exception:
    qdrant.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=models.VectorParams(
            size=1536,
            distance=models.Distance.COSINE
        ),
    )

# ── Embeddings ───────────────────────────────────────────
embeddings_model = OpenAIEmbeddings(model="text-embedding-3-small")

# ── Text splitter ────────────────────────────────────────
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    length_function=len,
    separators=["\n\n", "\n", ".", " ", ""],
)

# ── Tipuri acceptate ─────────────────────────────────────
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


# ────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────

def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[-1].lower()


def validate_file(filename: str, content_type: str) -> None:
    ext = get_file_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Format nesuportat: {ext}. Acceptăm PDF, DOCX și TXT.",
        )


async def extract_text(file: UploadFile) -> str:
    """Extrage textul din PDF, DOCX sau TXT."""
    content = await file.read()
    ext = get_file_extension(file.filename)

    try:
        if ext == ".pdf":
            return _extract_pdf(content)
        elif ext == ".docx":
            return _extract_docx(content)
        elif ext == ".txt":
            return _extract_txt(content)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Eroare la procesarea fișierului: {str(e)}",
        )


def _extract_pdf(content: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text


def _extract_docx(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Extragem și din tabele
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _extract_txt(content: bytes) -> str:
    # Încearcă UTF-8, fallback la latin-1
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="ignore")